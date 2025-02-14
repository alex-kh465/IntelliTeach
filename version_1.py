import streamlit as st
from io import BytesIO
from groq import Groq
import base64
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Initialize Groq API

import streamlit as st

API_KEY = st.secrets["api_keys"]["API_KEY"]
client = Groq(api_key=API_KEY)



# File uploader for Word documents
uploaded_file = st.file_uploader("Upload a Word Document", type=[".docx"])
# Extract text from uploaded .docx file
doc = Document(uploaded_file)
full_text = " ".join([para.text for para in doc.paragraphs])

# Display extracted text (for debugging)
st.subheader("Extracted Content")
st.write(full_text[:1000] + ("..." if len(full_text) > 1000 else ""))  # Show first 1000 chars


# Function to load the background image
def add_background_image(image_path):
    with open(image_path, "rb") as img_file:
        img_data = img_file.read()
        encoded_img = base64.b64encode(img_data).decode()
    
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-image: url('data:image/jpeg;base64,{encoded_img}');
            background-size: cover;
            background-position: center;
            height: 100vh;
        }}
        </style>
        """, unsafe_allow_html=True
    )

# Function to generate questions
def get_questions(api_key, model, subject, topic, num_questions, difficulty):
    client = Groq(api_key=api_key)
    
    try:
        if uploaded_file:
            prompt = (
                f"Create {num_questions} university-level multiple-choice questions based on the following content:\n{full_text} "
                f"categorized as '{difficulty}' difficulty.\n"
                "Each question must have 4 options (A, B, C, D) and a labeled 'Correct Answer'. "
                "Provide a summarized answer key at the end."
            )
        else:
            prompt = (
                f"Create {num_questions} university-level multiple-choice questions on the subject {subject} "
                f"focused on the topic '{topic}' categorized as '{difficulty}' difficulty.\n"
                "Each question must have 4 options (A, B, C, D) and a labeled 'Correct Answer'. "
                "Provide a summarized answer key at the end."
            )

        completion = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a helpful assistant who generates questions based on difficulty."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.7,
            max_tokens=2048,
            top_p=1,
            stream=True,
            stop=None,
        )

        questions = ""
        for chunk in completion:
            if hasattr(chunk.choices[0], "delta"):
                content = chunk.choices[0].delta.content or ""
                questions += content

        return questions
    except Exception as e:
        return f"Error: {e}"

# Function to generate DOCX file
def generate_docx(questions):
    doc = Document()
    doc.add_heading('Multiple Choice Questions', 0)

    doc.add_paragraph(questions)
    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)

    return docx_stream

# Function to generate PDF file
def generate_pdf(questions):
    pdf_buffer = BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    c.setFont("Helvetica", 10)
    
    text = c.beginText(40, 750)
    text.textLines(questions)
    
    c.drawText(text)
    c.showPage()
    c.save()

    pdf_buffer.seek(0)
    return pdf_buffer

# Download button
def download_button(content, filename, label):
    st.download_button(label, content, file_name=filename)

# Add background image
add_background_image("background.jpg")

# Streamlit App UI
st.title("University-Level MCQ Generator")
st.write("Generate randomized multiple-choice questions categorized by difficulty and optionally create a Google Form Quiz.")

# Use session state to persist data
for key in ["easy_questions", "medium_questions", "hard_questions"]:
    if key not in st.session_state:
        st.session_state[key] = ""

# Sidebar Inputs
st.sidebar.header("Configuration")
with st.sidebar:
    model = st.selectbox(
        "Select Model",
        ["llama3-8b-8192", "llama3-13b", "llama3-30b"],
        help="Choose the LLM for question generation.",
    )
    subject = st.text_input("Subject", help="Enter the subject name (e.g., Physics).")
    topic = st.text_input("Topic", help="Enter the topic name (e.g., Quantum Mechanics).")
    total_questions = st.number_input(
        "Total Questions", min_value=1, max_value=50, value=10, help="Total number of MCQs to generate."
    )

    st.markdown("### Difficulty Distribution")
    easy_pct = st.slider("Easy (%)", 0, 100, 30)
    medium_pct = st.slider("Medium (%)", 0, 100, 50)
    hard_pct = st.slider("Hard (%)", 0, 100, 20)

    # Validate percentage distribution
    if easy_pct + medium_pct + hard_pct != 100:
        st.error("The sum of percentages must equal 100.")

# Generate Questions per Section
num_easy = round((easy_pct / 100) * total_questions)
num_medium = round((medium_pct / 100) * total_questions)
num_hard = total_questions - num_easy - num_medium

if st.sidebar.button("Generate All Questions"):
    with st.spinner("Generating questions..."):
        st.session_state["easy_questions"] = get_questions(API_KEY, model, subject, topic, num_easy, "easy")
        st.session_state["medium_questions"] = get_questions(API_KEY, model, subject, topic, num_medium, "medium")
        st.session_state["hard_questions"] = get_questions(API_KEY, model, subject, topic, num_hard, "hard")
        st.success("Questions generated successfully!")

# Sections for Easy, Medium, and Hard Questions
for difficulty, key, num_questions in [("Easy", "easy_questions", num_easy), 
                                       ("Medium", "medium_questions", num_medium), 
                                       ("Hard", "hard_questions", num_hard)]:
    st.subheader(f"{difficulty} Questions")
    if st.session_state[key]:
        st.text_area(f"{difficulty} Questions", value=st.session_state[key], height=300)
    else:
        st.write(f"No {difficulty.lower()} questions generated yet.")
    
    # Buttons to regenerate questions independently
    if st.button(f"Regenerate {difficulty} Questions"):
        with st.spinner(f"Regenerating {difficulty.lower()} questions..."):
            st.session_state[key] = get_questions(API_KEY, model, subject, topic, num_questions, difficulty.lower())
            st.success(f"{difficulty} questions regenerated successfully!")
    
    if st.session_state[key]:
        download_button(st.session_state[key].encode('utf-8'), f"{difficulty.lower()}_questions.txt", f"Download {difficulty} Questions")

# Final Download for All Questions
if any(st.session_state[key] for key in ["easy_questions", "medium_questions", "hard_questions"]):
    st.subheader("Download Final Questions")
    final_questions = "\n\n".join([
        f"Easy Questions:\n{st.session_state['easy_questions']}" if st.session_state['easy_questions'] else "",
        f"Medium Questions:\n{st.session_state['medium_questions']}" if st.session_state['medium_questions'] else "",
        f"Hard Questions:\n{st.session_state['hard_questions']}" if st.session_state['hard_questions'] else ""
    ]).strip()

    # TXT download (final questions)
    download_button(final_questions.encode('utf-8'), "final_questions.txt", "Download Final Questions (TXT)")

    # DOCX download (final questions)
    docx_stream = generate_docx(final_questions)
    download_button(docx_stream, "final_questions.docx", "Download Final Questions (DOCX)")

    # PDF download (final questions)
    pdf_buffer = generate_pdf(final_questions)
    download_button(pdf_buffer, "final_questions.pdf", "Download Final Questions (PDF)")
